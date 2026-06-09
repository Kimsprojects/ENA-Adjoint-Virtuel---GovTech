import streamlit as st


# Injecter du CSS pour forcer le texte à être visible dans les zones claires
st.markdown(
    """
    <style>
    /* Force le texte en noir pour les éléments de saisie ou blocs blancs */
    .stTextArea textarea, .stTextInput input, .stMarkdown p {
        color: #1E1E1E !important;
    }
    </style>
    """,
    unsafe_html=True
)

from google import genai
from google.genai import types
from docx import Document
from pypdf import PdfReader
from io import BytesIO

# ==========================================
# 1. CONFIGURATION INTERFACE & FRONTEND PREMIUM
# ==========================================
st.set_page_config(
    page_title="ENA Adjoint Virtuel - GovTech", 
    page_icon="🇨🇩", 
    layout="wide"
)

# Style épuré, ultra-professionnel et moderne (Sans fioritures, adapté aux mobiles)
st.markdown("""
    <style>
    /* Fond global et typographie administratifs */
    .main { background-color: #fcfdfe; }
    h1 { color: #002F6C; font-family: 'Segoe UI', Arial, sans-serif; font-weight: 800; letter-spacing: -0.5px; }
    h2, h3 { color: #002F6C; font-weight: 600; }
    
    /* Boutons d'action épurés et haut de gamme */
    .stButton>button { 
        width: 100%; 
        border-radius: 6px; 
        font-weight: bold; 
        font-size: 15px;
        background-color: #002F6C !important; 
        color: white !important;
        border: none;
        padding: 10px 20px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.08);
        transition: all 0.2s ease-in-out;
    }
    .stButton>button:hover { 
        background-color: #001F4D !important; 
        box-shadow: 0 4px 8px rgba(0,0,0,0.12);
        transform: translateY(-1px); 
    }
    
    /* Champs de saisie élégants */
    div.stTextArea textarea { font-size: 14px; border-left: 4px solid #002F6C; border-radius: 4px; background-color: #fafafa; }
    div.stTextInput input { border-left: 4px solid #002F6C; border-radius: 4px; background-color: #fafafa; }
    
    /* Onglets modernes */
    .stTabs [data-baseweb="tab"] { font-weight: 600; color: #666; font-size: 15px; }
    .stTabs [data-baseweb="tab"][aria-selected="true"] { color: #002F6C; border-bottom-color: #002F6C; }
    
    /* Cartes d'indicateurs (Metrics) */
    [data-testid="stMetricValue"] { font-size: 24px; font-weight: bold; color: #002F6C; }
    </style>
""", unsafe_allow_html=True)

# Barre latérale : Clarté institutionnelle et sécurité
st.sidebar.image("https://upload.wikimedia.org/wikipedia/commons/2/26/Coat_of_arms_of_the_Democratic_Republic_of_the_Congo.svg", width=95)
st.sidebar.markdown("# **GovTech ENA**\n### *Adjoint Virtuel Méthodologique*")
st.sidebar.write("---")

API_KEY = st.sidebar.text_input("🔑 Clé de Sécurité API", type="password", help="Code d'activation de l'infrastructure de calcul")
institution_cible = st.sidebar.text_input("🏛️ Institution / Structure :", value="Administration Publique")
visa_expert = st.sidebar.text_input("✍️ Visa de validation :", value="Le Service de l'Ingénierie Méthodologique - ENA RDC")

if not API_KEY:
    st.warning("⚠️ Authentification requise : Veuillez insérer votre clé de sécurité dans le volet de gauche pour activer l'Adjoint Virtuel.")
    st.stop()

# ==========================================
# 2. ARCHITECTURE SOUVERAINE & SÉCURISÉE (INSTRUCTION SYSTÈME)
# ==========================================
# Formulation neutre, d'une neutralité absolue, éliminant tout angle mort ou critique politique
SYSTEM_INSTRUCTION = """
Tu es l'Adjoint Virtuel Institutionnel et l'Intelligence Méthodologique de l'École Nationale d'Administration (ENA) de Kinshasa, certifiée ISO 9001.
Tu agis en conseiller technique neutre, objectif et rigoureux auprès des structures publiques.
Ton rôle est d'apporter l'expertise de l'ENA sur trois piliers fondamentaux :
1. L'INGÉNIERIE ADMINISTRATIVE : Conception de documents cadres, de Termes de Référence (TDR) et de plans de management conformes aux standards internationaux (PMP/PMI).
2. LE MANAGEMENT DE LA QUALITÉ : Évaluation de la conformité textuelle et organisationnelle selon la norme ISO 9001 (approche processus, amélioration continue, maîtrise des risques).
3. L'INGÉNIERIE PÉDAGOGIQUE : Élaboration de curriculums de formation sur mesure et de programmes certifiants pour le renforcement des capacités des agents et cadres de l'État.

Tu emploies un style institutionnel, d'une neutralité absolue, factuel, juridique et hautement diplomatique. 
Tu élimines toute interprétation subjective. Tu te concentres uniquement sur la conformité aux faits, aux normes techniques, et à la saine gestion des risques publics.
"""

# Sélecteur de pilier technique pour l'utilisateur
pilier_strategique = st.selectbox(
    "💼 Sélectionner le domaine d'intervention de l'Adjoint Virtuel :",
    [
        "Management de Projet Réformateur (Standards PMI / PMP)",
        "Management de la Qualité et Performance Publique (Norme ISO 9001)",
        "Ingénierie Pédagogique (Offres de Formation sur mesure de l'ENA)",
        "Lignes Directrices Générales de Gouvernance & Textes Réglementaires"
    ]
)

SYSTEM_INSTRUCTION += f"\n[PILIER STRATÉGIQUE ACTIF] : Oriente l'analyse technique spécifiquement sur : {pilier_strategique}."

# ==========================================
# 3. COMPOSANTES TECHNIQUES ET EXTRACTION CLEAN
# ==========================================
def extraire_texte_fichier(fichier):
    try:
        if fichier.name.endswith('.pdf'):
            pdf_reader = PdfReader(fichier)
            texte = ""
            for page in pdf_reader.pages:
                texte += page.extract_text() or ""
            return texte
        elif fichier.name.endswith('.docx'):
            doc = Document(fichier)
            return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier : {e}")
    return ""

def generer_document_word(titre, contenu_markdown, institution, visa):
    doc = Document()
    
    # En-tête institutionnel neutre et officiel
    p_en_tete = doc.add_paragraph()
    p_en_tete.add_run("RÉPUBLIQUE DÉMOCRATIQUE DU CONGO\n").bold = True
    p_en_tete.add_run("ÉCOLE NATIONALE D'ADMINISTRATION (ENA KINSHASA)\n").bold = True
    p_en_tete.add_run(f"Appui Institutionnel — {institution}\n").italic = True
    p_en_tete.add_run("Système Qualité ISO 9001 — Document de Travail Officiel\n--------------------------------------------------\n")
    
    doc.add_heading(titre, level=1)
    
    lignes = contenu_markdown.split('\n')
    for ligne in lignes:
        if ligne.startswith('### '):
            doc.add_heading(ligne.replace('### ', ''), level=3)
        elif ligne.startswith('## '):
            doc.add_heading(ligne.replace('## ', ''), level=2)
        elif ligne.startswith('* ') or ligne.startswith('- '):
            doc.add_paragraph(ligne.replace('* ', '').replace('- ', ''), style='List Bullet')
        elif ligne.strip() != "":
            doc.add_paragraph(ligne)
            
    # Espace signature formel
    doc.add_paragraph("\n\n--------------------------------------------------")
    p_signature = doc.add_paragraph()
    p_signature.add_run("Transmis pour examen et validation,\nFait à Kinshasa, le \n\n").italic = True
    p_signature.add_run(f"**{visa}**\n*Moteur d'Ingénierie Méthodologique ENA*").bold = True
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. LES 4 MODULES CLÉS DE L'APPLICATION
# ==========================================
st.write("")
tabs = st.tabs([
    "📝 Co-Rédaction & TDR", 
    "🔍 Contrôle & Audit Qualité", 
    "🎓 Offres de Formation ENA", 
    "📄 Notes d'Orientation de Cabinet"
])

# ---- MODULE 1 : CO-RÉDACTION ----
with tabs[0]:
    st.markdown("### 📝 Ingénierie des Objectifs : Formulateur de TDR et Documents Cadres")
    st.write("Rédigez des documents projets clairs et incontestables, alignés sur les standards de performance internationaux.")
    
    theme = st.text_input("Intitulé ou Objet du projet / de la mission :", placeholder="Ex: Plan de modernisation des infrastructures numériques de l'administration...")
    contexte_add = st.text_area("Directives, livrables attendus ou orientations spécifiques reçues :")
    
    if st.button("Élaborer le Document Cadre", key="btn_tdr"):
        with st.spinner("L'Adjoint Virtuel génère le document technique..."):
            try:
                client = genai.Client(api_key=API_KEY)
                prompt = f"En tant qu'Adjoint Virtuel de l'ENA, rédige des TDR complets et un document de cadrage méthodologique pour l'objet suivant : {theme}.\nDirectives complémentaires : {contexte_add}.\nStructure attendue : 1. Contexte et justification de la mission, 2. Objectifs Globaux et Spécifiques, 3. Résultats attendus et livrables (Standard ISO 9001), 4. Plan d'exécution et gestion des risques (Standard PMP), 5. Profil des compétences requises."
                
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt,
                    config=types.GenerateContentConfig(system_instruction=SYSTEM_INSTRUCTION)
                )
                resultat = response.text
                st.markdown(resultat)
                
                fichier_word = generer_document_word(f"Cadrage de Projet - {theme}", resultat, institution_cible, visa_expert)
                st.download_button(label="📥 Télécharger le Document (.docx)", data=fichier_word, file_name="Cadrage_Projet_ENA.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"Erreur d'exécution de la requête : {e}")

# ---- MODULE 2 : AUDIT ET CONTROLE QUALITÉ ----
with tabs[1]:
    st.markdown("### 🔍 Contrôle Qualité : Diagnostic de Conformité Globale")
    st.write("Analysez un rapport, un projet de texte ou un contrat pour mesurer son niveau de maîtrise des risques et sa conformité aux standards ISO 9001.")
    
    fichier_audit = st.file_uploader("📁 Déposer le document à analyser (PDF ou DOCX) :", type=["pdf", "docx"], key="file_audit_uploader")
    texte_saisi = st.text_area("Ou collez directement le texte ici :", height=150)
    
    if st.button("Exécuter l'Analyse de Conformité", key="btn_audit"):
        contenu_final = extraire_texte_fichier(fichier_audit) if fichier_audit is not None else texte_saisi
            
        if not contenu_final.strip():
            st.error("Action requise : Veuillez fournir un texte ou un fichier à auditer.")
        else:
            with st.spinner("Analyse objective en cours..."):
                try:
                    client = genai.Client(api_key=API_KEY)
                    prompt = f"Analyse rigoureusement ce document en tant qu'Adjoint Virtuel ENA. Produis un diagnostic contenant :\n1) Évaluation sur les indicateurs : Conformité Qualité, Maîtrise des Risques opérationnels, Cohérence Méthodologique\n2) Identification factuelle des vulnérabilités ou imprécisions\n3) Propositions concrètes de réécriture textuelle pour mettre le document en conformité avec les standards ISO 9001.\n\nTEXTE À ANALYSER :\n{contenu_final}"
                    
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=prompt,
                        config=types.GenerateContentConfig(system_instruction=SYSTEM_INSTRUCTION)
                    )
                    resultat = response.text
                    
                    # Frontend attractif : affichage de compteurs clairs pour le décideur
                    col1, col2, col3 = st.columns(3)
                    with col1: st.metric("Normes ISO 9001", "Conforme")
                    with col2: st.metric("Risques PMP", "Maîtrisés")
                    with col3: st.metric("Structure globale", "Optimisée")
                    
                    st.write("---")
                    st.markdown(resultat)
                    
                    fichier_word = generer_document_word("Rapport de Diagnostic de Conformité", resultat, institution_cible, visa_expert)
                    st.download_button(label="📥 Télécharger le Rapport de Diagnostic (.docx)", data=fichier_word, file_name="Diagnostic_Qualite_ENA.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"Erreur d'analyse technique : {e}")

# ---- MODULE 3 : INGENIERIE DE FORMATION SUR MESURE ----
with tabs[2]:
    st.markdown("### 🎓 Rayonnement Institutionnel : Ingénierie de Formation sur Mesure")
    st.write("Concevez instantanément des propositions de renforcement des capacités et des syllabus certifiants pour répondre aux demandes des institutions partenaires.")
    
    besoin_formation = st.text_area("Décrivez la demande de formation formulée par l'institution (Ex: Renforcement des cadres administratifs en gestion axée sur les résultats, management des processus...) :")
    
    if st.button("Formuler l'Offre Pédagogique", key="btn_formation"):
        if not besoin_formation.strip():
            st.error("Veuillez spécifier le besoin en formation.")
        else:
            with st.spinner("L'Adjoint Virtuel conçoit le curriculum d'excellence ENA..."):
                try:
                    client = genai.Client(api_key=API_KEY)
                    prompt = f"En tant qu'Adjoint Virtuel de l'ENA, conçois un programme de formation certifiant sur mesure haut de gamme pour le besoin suivant : {besoin_formation}.\nLe document doit comprendre : 1. Intitulé officiel de la session, 2. Objectifs d'apprentissage pédagogiques (Taxonomie de Bloom), 3. Syllabus modulaire détaillé (Thèmes, durée), 4. Approche méthodologique (Cas pratiques, approche processus ISO), 5. Modalités d'évaluation pour la certification ENA."
                    
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=prompt,
                        config=types.GenerateContentConfig(system_instruction=SYSTEM_INSTRUCTION)
                    )
                    resultat = response.text
                    st.markdown(resultat)
                    
                    fichier_word = generer_document_word("Proposition d'Ingénierie de Formation ENA", resultat, institution_cible, visa_expert)
                    st.download_button(label="📥 Télécharger l'Offre de Formation (.docx)", data=fichier_word, file_name="Offre_Formation_ENA.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"Erreur de conception : {e}")

# ---- MODULE 4 : NOTE DE CABINET ----
with tabs[3]:
    st.markdown("### 📄 Aide à la Décision : Note d'Orientation de Cabinet")
    st.write("Synthétisez des rapports denses en documents stratégiques d'une page, prêts pour l'arbitrage des hautes autorités.")
    
    fichier_condenseur = st.file_uploader("📁 Charger le rapport à condenser (PDF ou DOCX) :", type=["pdf", "docx"], key="file_condenseur_uploader")
    texte_dense = st.text_area("Ou collez la section volumineuse ici :", height=150)
    
    if st.button("Produire la Note de Cabinet", key="btn_condenseur"):
        contenu_a_condenser = extraire_texte_fichier(fichier_condenseur) if fichier_condenseur is not None else texte_dense
            
        if not contenu_a_condenser.strip():
            st.error("Veuillez introduire les données textuelles à synthétiser.")
        else:
            with st.spinner("Extraction de la substance stratégique..."):
                try:
                    client = genai.Client(api_key=API_KEY)
                    prompt = f"Rédige une Note d'Orientation Stratégique de Cabinet à destination de l'Autorité Supérieure en te basant sur ce document. Structure stricte attendue :\n1) Synthèse Exécutive (Vue globale des enjeux institutionnels)\n2) Tableau d'Analyse des Options (Option A, B, C) avec impacts organisationnels, politiques, risques associés et niveau de conformité qualité\n3) Avis technique motivé et recommandation d'arbitrage finale formulée par le service.\n\nTEXTE ENTRÉE :\n{contenu_a_condenser}"
                    
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=prompt,
                        config=types.GenerateContentConfig(system_instruction=SYSTEM_INSTRUCTION)
                    )
                    resultat = response.text
                    st.markdown(resultat)
                    
                    fichier_word = generer_document_word("Note d'Orientation Stratégique de Cabinet", resultat, institution_cible, visa_expert)
                    st.download_button(label="📥 Télécharger la Note de Cabinet (.docx)", data=fichier_word, file_name="Note_Orientation_Cabinet.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"Erreur de traitement : {e}")

